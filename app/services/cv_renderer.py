"""CV Renderer — produces PDF bytes and JSON dict from a CVData model."""
from __future__ import annotations

import logging
from io import BytesIO
from pathlib import Path

from app.schemas import CVData

logger = logging.getLogger(__name__)

_TEMPLATE_DIR = Path(__file__).parent / "cv_export" / "templates" / "cv"

# A4 page height in CSS pixels at 96 DPI: 297mm / 25.4 * 96 ≈ 1122.5
_A4_PAGE_HEIGHT_PX = 1123


def render_cv_html(cv: CVData, template_name: str = "default", **context) -> str:
    """Render CV data to an HTML string using Jinja2."""
    from jinja2 import Environment, FileSystemLoader

    env = Environment(loader=FileSystemLoader(str(_TEMPLATE_DIR)), autoescape=True)
    template = env.get_template(f"{template_name}.html")
    return template.render(cv=cv, **context)


def render_cv_pdf(cv: CVData, template_name: str = "default") -> bytes:
    """Render CV to PDF bytes via Playwright (headless Chromium)."""
    from playwright.sync_api import sync_playwright

    html_str = render_cv_html(cv, template_name)
    with sync_playwright() as p:
        browser = p.chromium.launch()
        page = browser.new_page()
        page.set_content(html_str, wait_until="load")
        pdf_bytes = page.pdf(format="A4", print_background=True)
        browser.close()
    return pdf_bytes


def render_cv_json(cv: CVData) -> dict:
    """Serialise CVData to a plain dict with ISO-formatted dates."""
    return cv.model_dump()


def render_tailored_pdf(cv: CVData) -> bytes:
    """Render the tailored CV template to PDF bytes.

    Fit cascade (see memory/guide.md §13): the CVData itself is already pruned
    upstream (cut content, tightened language). Here we only handle the final
    typographic tightening — if the normal-density render overflows A4, we
    re-render with compact mode (smaller body, tighter line-height and gaps).
    """
    from playwright.sync_api import sync_playwright

    def _render(compact: bool) -> tuple[bytes, int]:
        html_str = render_cv_html(cv, template_name="tailored", compact=compact)
        with sync_playwright() as p:
            browser = p.chromium.launch()
            page = browser.new_page(viewport={"width": 794, "height": _A4_PAGE_HEIGHT_PX})
            page.set_content(html_str, wait_until="load")
            content_height = page.evaluate("document.documentElement.scrollHeight")
            pdf_bytes = page.pdf(
                format="A4",
                print_background=True,
                margin={"top": "0", "bottom": "0", "left": "0", "right": "0"},
            )
            browser.close()
        return pdf_bytes, int(content_height)

    pdf_bytes, height_px = _render(compact=False)
    if height_px > _A4_PAGE_HEIGHT_PX:
        logger.info(
            "Tailored CV overflowed A4 (%dpx > %dpx); re-rendering in compact mode",
            height_px, _A4_PAGE_HEIGHT_PX,
        )
        pdf_bytes, _ = _render(compact=True)
    return pdf_bytes


def render_cover_letter_pdf(text: str, title: str = "", company: str = "") -> bytes:
    """Render plain-text cover letter to PDF bytes via Playwright."""
    from playwright.sync_api import sync_playwright
    import html as _html

    heading = ""
    if title or company:
        parts = [p for p in [title, company] if p]
        heading = f"<h2 style='margin:0 0 20px;font-size:13px;font-weight:600;color:#444;font-family:Georgia,serif'>{_html.escape(' — '.join(parts))}</h2>"

    # Split on paragraph breaks; within each paragraph preserve line breaks via <br>
    paragraphs = text.strip().split("\n\n")
    paras_html = "".join(
        f"<p>{'<br>'.join(_html.escape(line) for line in para.splitlines())}</p>"
        for para in paragraphs
        if para.strip()
    )

    html_str = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
  * {{ box-sizing: border-box; }}
  body {{
    font-family: Georgia, serif;
    font-size: 12pt;
    line-height: 1.75;
    color: #222;
    margin: 0;
    padding: 0;
  }}
  p {{ margin: 0 0 14px; }}
</style>
</head><body>
{heading}
{paras_html}
</body></html>"""

    with sync_playwright() as p:
        browser = p.chromium.launch()
        page = browser.new_page()
        page.set_content(html_str, wait_until="load")
        pdf_bytes = page.pdf(
            format="A4",
            print_background=True,
            margin={"top": "25mm", "bottom": "25mm", "left": "25mm", "right": "25mm"},
        )
        browser.close()
    return pdf_bytes


def render_cover_letter_docx(text: str, title: str = "", company: str = "") -> bytes:
    """Render plain-text cover letter to DOCX bytes via python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, Mm
    except ImportError as exc:
        raise ImportError("python-docx is required. Run: pip install python-docx") from exc

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)

    normal = doc.styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(11)

    if title or company:
        parts = [p for p in [title, company] if p]
        h = doc.add_paragraph()
        h.paragraph_format.space_after = Pt(14)
        run = h.add_run(" — ".join(parts))
        _style_run(run, size=12, bold=True, color=_INK_700)

    for para in text.strip().split("\n\n"):
        para = para.strip()
        if not para:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        for i, line in enumerate(para.splitlines()):
            if i > 0:
                p.add_run().add_break()
            run = p.add_run(line)
            _style_run(run, size=11, color=_INK_700)

    from io import BytesIO
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _bullets_from_description(description: str | None, limit: int = 5) -> list[str]:
    if not description:
        return []
    out: list[str] = []
    for line in description.splitlines():
        clean = line.strip().lstrip("•").lstrip("-").lstrip("·").strip()
        if clean:
            out.append(clean)
        if len(out) >= limit:
            break
    return out


# ── DOCX helpers ──────────────────────────────────────────────────────────────

_INK_900 = (0x11, 0x11, 0x11)
_INK_700 = (0x33, 0x33, 0x33)
_INK_500 = (0x66, 0x66, 0x66)
_ACCENT = (0x2D, 0x2D, 0x2D)


def _style_run(run, size=None, bold=None, color=None, italic=None, font_name=None):
    """Apply common run properties; None means leave unchanged."""
    from docx.shared import Pt, RGBColor

    if font_name is not None:
        run.font.name = font_name
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def _set_char_spacing_twips(run, twips: int) -> None:
    """Emulate CSS letter-spacing via Word's <w:spacing w:val="..."/> (in twips)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    rpr = run._element.get_or_add_rPr()
    spacing = rpr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        rpr.append(spacing)
    spacing.set(qn("w:val"), str(twips))


def _add_section_heading(doc, text: str) -> None:
    from docx.shared import Pt

    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(4)
    run = p.add_run(text.upper())
    _style_run(run, size=11, bold=True, color=_ACCENT)
    # 0.08em at 11pt ≈ 0.88pt → ~18 twips (1pt = 20 twips, so use 16 for safety)
    _set_char_spacing_twips(run, 16)


def render_tailored_docx(cv: CVData) -> bytes:
    """Render the tailored CV as a DOCX byte stream via python-docx.

    Mirrors the HTML/PDF design: single font family, ink color tokens,
    uppercase letter-spaced section headings, real Word bullet lists.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Mm
    except ImportError as exc:
        raise ImportError("python-docx is required for DOCX rendering. Run: pip install python-docx") from exc

    doc = Document()

    # Page margins — 18mm all sides per guide §1.
    section = doc.sections[0]
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)

    # Default body font. Calibri ships with Word; Inter is preferred when present
    # on the host, but most Word users don't have it, so default to Calibri to
    # avoid font-substitution warnings.
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    # Header — name
    if cv.full_name:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(cv.full_name)
        _style_run(run, size=24, bold=True, color=_INK_900)

    if cv.headline:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(cv.headline)
        _style_run(run, size=11, color=_INK_500)

    contact_parts = [x for x in (cv.location, cv.email, cv.phone, cv.profile_url) if x]
    if contact_parts:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run("  ·  ".join(contact_parts))
        _style_run(run, size=9.5, color=_INK_700)

    if cv.tailored_summary:
        _add_section_heading(doc, "Summary")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(cv.tailored_summary)
        _style_run(run, size=10, color=_INK_700)

    if cv.prioritized_skills:
        _add_section_heading(doc, "Skills")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(", ".join(cv.prioritized_skills))
        _style_run(run, size=10, color=_INK_700)

    if cv.experience:
        _add_section_heading(doc, "Experience")
        for exp in cv.experience:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            title_run = p.add_run(exp.title or "")
            _style_run(title_run, size=10.5, bold=True, color=_INK_900)
            if exp.company:
                company_run = p.add_run(f"  —  {exp.company}")
                _style_run(company_run, size=10.5, color=_INK_700)

            date_bits = []
            if exp.start_date:
                date_bits.append(exp.start_date)
            if exp.is_current:
                date_bits.append("Present")
            elif exp.end_date:
                date_bits.append(exp.end_date)
            date_text = " – ".join(date_bits)
            if exp.location:
                date_text = f"{date_text}  ·  {exp.location}" if date_text else exp.location
            if date_text:
                meta_run = p.add_run(f"    {date_text}")
                _style_run(meta_run, size=9.5, color=_INK_500)

            for bullet in _bullets_from_description(exp.description, limit=5):
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.space_before = Pt(0)
                bp.paragraph_format.space_after = Pt(2)
                brun = bp.add_run(bullet)
                _style_run(brun, size=10, color=_INK_700)

    if cv.education:
        _add_section_heading(doc, "Education")
        for edu in cv.education:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(edu.school or "")
            _style_run(run, size=10.5, bold=True, color=_INK_900)
            extras = []
            if edu.degree:
                extras.append(edu.degree)
            if edu.field_of_study:
                extras.append(edu.field_of_study)
            if extras:
                ex_run = p.add_run("  —  " + ", ".join(extras))
                _style_run(ex_run, size=10.5, color=_INK_700)
            years = " – ".join([y for y in (edu.start_year, edu.end_year) if y])
            if years:
                meta_run = p.add_run(f"    {years}")
                _style_run(meta_run, size=9.5, color=_INK_500)

    if cv.certifications:
        _add_section_heading(doc, "Certifications")
        for c in cv.certifications:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            name_run = p.add_run(c.name or "")
            _style_run(name_run, size=10, color=_INK_700)
            if c.issuing_org:
                org_run = p.add_run(f" — {c.issuing_org}")
                _style_run(org_run, size=10, color=_INK_500)
            if c.issue_date:
                date_run = p.add_run(f" ({c.issue_date})")
                _style_run(date_run, size=10, color=_INK_500)

    if cv.projects:
        _add_section_heading(doc, "Projects")
        for proj in cv.projects:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(proj.name or "")
            _style_run(run, size=10.5, bold=True, color=_INK_900)
            if proj.description:
                dp = doc.add_paragraph()
                dp.paragraph_format.space_after = Pt(2)
                drun = dp.add_run(proj.description)
                _style_run(drun, size=10, color=_INK_700)
            if proj.url:
                up = doc.add_paragraph()
                up.paragraph_format.space_after = Pt(2)
                urun = up.add_run(proj.url)
                _style_run(urun, size=9.5, color=_INK_500)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
